"""Utilities for exporting requirements into multiple formats."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, UTC
from io import BytesIO
from pathlib import Path
from collections.abc import Iterable, Mapping, Sequence
import json
import re

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, StyleSheet1, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape as xml_escape

import markdown
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from ..i18n import _
from ..util.time import format_datetime_for_humans
from .document_store import (
    Document,
    DocumentNotFoundError,
    get_document_revision,
    load_documents,
    load_requirements,
)
from .document_store import label_color as resolved_label_color
from .markdown_utils import (
    convert_markdown_math,
    normalize_escaped_newlines,
    sanitize_html,
    strip_markdown,
)
from .model import Requirement

__all__ = [
    "DocumentExport",
    "RequirementExport",
    "RequirementExportLink",
    "RequirementExportView",
    "build_requirement_export",
    "build_requirement_export_from_requirements",
    "export_card_field_order",
    "render_requirements_html",
    "render_requirements_markdown",
    "render_requirements_docx",
    "render_requirements_pdf",
]



def _resolve_field_content(
    value: str | None,
    *,
    empty_field_placeholder: str | None,
) -> str | None:
    if value:
        return value
    if empty_field_placeholder is not None:
        return empty_field_placeholder
    return None


@dataclass(slots=True)
class RequirementExportLink:
    """Representation of a requirement relationship for export."""

    rid: str
    title: str | None
    exists: bool
    suspect: bool


@dataclass(slots=True)
class RequirementExportView:
    """View model combining requirement metadata and link summaries."""

    requirement: Requirement
    document: Document
    links: list[RequirementExportLink]


@dataclass(slots=True)
class DocumentExport:
    """Collection of requirements grouped by their document."""

    document: Document
    requirements: list[RequirementExportView]


@dataclass(slots=True)
class RequirementExport:
    """High-level container for exported data."""

    documents: list[DocumentExport]
    selected_prefixes: tuple[str, ...]
    generated_at: datetime
    base_path: Path


@dataclass(slots=True)
class RequirementLinkPreview:
    """Compact preview payload for interactive trace links in HTML export."""

    rid: str
    title: str
    status: str
    req_type: str
    statement_preview: str
    exists: bool = True
    suspect: bool = False


def _normalize_prefixes(prefixes: Sequence[str] | None, docs: Mapping[str, Document]) -> tuple[str, ...]:
    if prefixes is None:
        return tuple(sorted(docs))
    order: list[str] = []
    seen: set[str] = set()
    for prefix in prefixes:
        if prefix not in docs:
            raise DocumentNotFoundError(prefix)
        if prefix in seen:
            continue
        seen.add(prefix)
        order.append(prefix)
    return tuple(order)


def build_requirement_export(
    root: str | Path,
    *,
    prefixes: Sequence[str] | None = None,
) -> RequirementExport:
    """Load requirements and assemble a deterministic export representation."""
    root_path = Path(root)
    docs = load_documents(root_path)
    if not docs and not root_path.is_dir():
        raise FileNotFoundError(root_path)
    ordered_prefixes = _normalize_prefixes(prefixes, docs)
    requirements = load_requirements(root_path, prefixes=ordered_prefixes or None, docs=docs)
    by_rid = {req.rid: req for req in requirements}

    grouped: dict[str, DocumentExport] = {
        prefix: DocumentExport(document=docs[prefix], requirements=[])
        for prefix in ordered_prefixes
    }
    for req in requirements:
        export_doc = grouped[req.doc_prefix]
        links: list[RequirementExportLink] = []
        for link in req.links:
            target = by_rid.get(link.rid)
            links.append(
                RequirementExportLink(
                    rid=link.rid,
                    title=target.title if target else None,
                    exists=bool(target),
                    suspect=getattr(link, "suspect", False),
                )
            )
        export_doc.requirements.append(
            RequirementExportView(requirement=req, document=export_doc.document, links=links)
        )

    ordered_documents = [grouped[prefix] for prefix in ordered_prefixes]
    return RequirementExport(
        documents=ordered_documents,
        selected_prefixes=ordered_prefixes,
        generated_at=datetime.now(UTC),
        base_path=root_path,
    )


def build_requirement_export_from_requirements(
    requirements: Sequence[Requirement],
    docs: Mapping[str, Document],
    *,
    base_path: Path,
    prefixes: Sequence[str] | None = None,
    link_lookup: Sequence[Requirement] | None = None,
) -> RequirementExport:
    """Build export view model using preloaded requirements."""
    if prefixes is None:
        ordered_prefixes = tuple(sorted({req.doc_prefix for req in requirements}))
    else:
        ordered_prefixes = _normalize_prefixes(prefixes, docs)
    for prefix in ordered_prefixes:
        if prefix not in docs:
            raise DocumentNotFoundError(prefix)

    grouped: dict[str, DocumentExport] = {
        prefix: DocumentExport(document=docs[prefix], requirements=[])
        for prefix in ordered_prefixes
    }
    link_source = requirements if link_lookup is None else link_lookup
    by_rid = {req.rid: req for req in link_source}
    for req in requirements:
        if req.doc_prefix not in grouped:
            continue
        export_doc = grouped[req.doc_prefix]
        links: list[RequirementExportLink] = []
        for link in req.links:
            target = by_rid.get(link.rid)
            links.append(
                RequirementExportLink(
                    rid=link.rid,
                    title=target.title if target else None,
                    exists=bool(target),
                    suspect=getattr(link, "suspect", False),
                )
            )
        export_doc.requirements.append(
            RequirementExportView(requirement=req, document=export_doc.document, links=links)
        )

    ordered_documents = [grouped[prefix] for prefix in ordered_prefixes]
    return RequirementExport(
        documents=ordered_documents,
        selected_prefixes=ordered_prefixes,
        generated_at=datetime.now(UTC),
        base_path=base_path,
    )




def _normalize_export_fields(fields: Iterable[str] | None) -> set[str] | None:
    if fields is None:
        return None
    return {field for field in fields if field}


def _should_render_field(selected_fields: set[str] | None, field: str) -> bool:
    if selected_fields is None:
        return True
    return field in selected_fields


def _localize_enum_code(value: str | None) -> str | None:
    if not value:
        return value
    msgid = value.replace("_", " ").strip().capitalize()
    if not msgid:
        return value
    return _(msgid)


def _requirement_heading(req: Requirement, selected_fields: set[str] | None) -> str:
    if _should_render_field(selected_fields, "title"):
        return f"{req.rid} — {req.title or _('(no title)')}"
    return req.rid


_EXPORT_META_FIELDS: tuple[tuple[str, str, bool], ...] = (
    ("type", "Requirement type", True),
    ("status", "Status", True),
    ("priority", "Priority", True),
    ("owner", "Owner", False),
    ("labels", "Labels", False),
    ("source", "Source", False),
    ("modified_at", "Modified at", False),
    ("approved_at", "Approved at", False),
    ("revision", "Revision", False),
)

_EXPORT_SECTION_FIELDS: tuple[tuple[str, str], ...] = (
    ("statement", "Requirement text"),
    ("acceptance", "Acceptance criteria"),
    ("conditions", "Conditions"),
    ("rationale", "Rationale"),
    ("assumptions", "Assumptions"),
    ("notes", "Notes"),
)


def export_card_field_order() -> tuple[str, ...]:
    """Return field identifiers in the order used by card-like exports."""
    meta_fields = (field for field, _label, _use_code in _EXPORT_META_FIELDS)
    section_fields = (field for field, _label in _EXPORT_SECTION_FIELDS)
    return ("rid", "title", *meta_fields, *section_fields)


def _meta_field_value(req: Requirement, field: str) -> str | None:
    if field == "type":
        return _localize_enum_code(req.type.value)
    if field == "status":
        return _localize_enum_code(req.status.value)
    if field == "priority":
        return _localize_enum_code(getattr(req.priority, "value", None))
    if field == "owner":
        return req.owner or None
    if field == "labels":
        return ", ".join(sorted(req.labels)) if req.labels else None
    if field == "source":
        return req.source or None
    if field == "modified_at":
        return req.modified_at or None
    if field == "approved_at":
        return req.approved_at or None
    if field == "revision":
        return str(req.revision)
    return None


def _section_field_value(req: Requirement, field: str) -> str | None:
    if field == "statement":
        return req.statement
    if field == "acceptance":
        return req.acceptance or ""
    if field == "conditions":
        return req.conditions
    if field == "rationale":
        return req.rationale
    if field == "assumptions":
        return req.assumptions
    if field == "notes":
        return req.notes
    return None


def _normalized_labels(req: Requirement) -> tuple[str, ...]:
    labels: list[str] = []
    seen: set[str] = set()
    for raw in req.labels:
        cleaned = raw.strip()
        if not cleaned:
            continue
        key = cleaned.casefold()
        if key in seen:
            continue
        seen.add(key)
        labels.append(cleaned)
    return tuple(sorted(labels, key=str.casefold))


def _group_requirement_views_by_labels(
    views: Sequence[RequirementExportView],
    *,
    unlabeled_title: str,
    label_group_mode: str,
) -> list[tuple[str, list[RequirementExportView]]]:
    groups: list[tuple[str, list[RequirementExportView]]] = []
    lookup: dict[str, list[RequirementExportView]] = {}

    def _append(group_title: str, view: RequirementExportView) -> None:
        bucket = lookup.get(group_title)
        if bucket is None:
            bucket = []
            lookup[group_title] = bucket
            groups.append((group_title, bucket))
        bucket.append(view)

    for view in views:
        labels = _normalized_labels(view.requirement)
        if not labels:
            _append(unlabeled_title, view)
            continue
        if label_group_mode == "label_set":
            _append(", ".join(labels), view)
            continue
        for label in labels:
            _append(label, view)
    return groups


def _label_palette(export: RequirementExport) -> dict[str, str]:
    palette: dict[str, str] = {}
    for doc_export in export.documents:
        for label in doc_export.document.labels.defs:
            palette[label.key.casefold()] = resolved_label_color(label)
    return palette


def _collect_used_label_rows(export: RequirementExport) -> list[tuple[str, str, str | None]]:
    palette = _label_palette(export)
    label_titles: dict[str, str] = {}
    for doc_export in export.documents:
        for label in doc_export.document.labels.defs:
            if label.key not in label_titles:
                label_titles[label.key] = label.title
    used_labels: set[str] = set()
    for doc_export in export.documents:
        for view in doc_export.requirements:
            used_labels.update(_normalized_labels(view.requirement))
    rows: list[tuple[str, str, str | None]] = []
    for label in sorted(used_labels, key=str.casefold):
        rows.append((label, label_titles.get(label, ""), palette.get(label.casefold())))
    return rows


def _normalize_hex_color(color: str | None) -> str | None:
    if not color:
        return None
    value = color.strip()
    if len(value) == 7 and value.startswith("#"):
        digits = value[1:]
        if all(ch in "0123456789abcdefABCDEF" for ch in digits):
            return digits.upper()
    return None


def _text_color_for_background(hex_color: str) -> str:
    red = int(hex_color[0:2], 16)
    green = int(hex_color[2:4], 16)
    blue = int(hex_color[4:6], 16)
    luminance = (0.299 * red) + (0.587 * green) + (0.114 * blue)
    return "000000" if luminance >= 150 else "FFFFFF"


def _render_html_label_chip(label: str, color: str | None) -> str:
    background = _normalize_hex_color(color)
    if not background:
        return f"<span class='label-chip'>{_escape_html(label)}</span>"
    text_color = _text_color_for_background(background)
    return (
        f"<span class='label-chip' style='background:#{background};color:#{text_color};border-color:#{background};'>"
        f"{_escape_html(label)}</span>"
    )

def _format_markdown_table_cell(text: str) -> str:
    normalized = text.strip("\n")
    if not normalized:
        return ""
    normalized = normalized.replace("|", "\\|")
    return "<br>".join(normalized.splitlines())


def _render_generated_at(export: RequirementExport) -> str:
    return format_datetime_for_humans(export.generated_at)


def _document_revision_label(document: Document) -> str:
    revision = get_document_revision(document)
    return _("rev {revision}").format(revision=revision)


def _export_revisions_summary(export: RequirementExport) -> str:
    parts = [
        f"{doc_export.document.prefix} {_document_revision_label(doc_export.document)}"
        for doc_export in export.documents
    ]
    return ", ".join(parts)


def render_requirements_markdown(
    export: RequirementExport,
    *,
    title: str | None = None,
    empty_field_placeholder: str | None = None,
    fields: Iterable[str] | None = None,
    group_by_labels: bool = False,
    unlabeled_group_title: str | None = None,
    label_group_mode: str = "per_label",
) -> str:
    """Render export data as Markdown."""
    selected_fields = _normalize_export_fields(fields)
    heading = title or _('Requirements export')
    parts: list[str] = [f"# {heading}", ""]
    parts.append(
        f"_{_('Generated at')} {_render_generated_at(export)} {_('for documents')}: {', '.join(export.selected_prefixes)}._"
    )
    parts.append(f"_{_('Document revisions')}: {_export_revisions_summary(export)}._")
    parts.append("")
    if _should_render_field(selected_fields, "labels"):
        label_rows = _collect_used_label_rows(export)
        if label_rows:
            parts.append(f"## {_('Labels')}")
            parts.append("")
            parts.append("| | |")
            parts.append("| --- | --- |")
            for label, description, _color in label_rows:
                value = _format_markdown_table_cell(description)
                parts.append(f"| {label} | {value} |")
            parts.append("")

    for doc in export.documents:
        parts.append(
            f"## {doc.document.title} ({doc.document.prefix}, {_document_revision_label(doc.document)})"
        )
        parts.append("")
        if group_by_labels:
            group_iter = _group_requirement_views_by_labels(
                doc.requirements,
                unlabeled_title=unlabeled_group_title or _('Without labels'),
                label_group_mode=label_group_mode,
            )
        else:
            group_iter = [("", list(doc.requirements))]

        for group_title, group_views in group_iter:
            heading_level = "###"
            if group_by_labels:
                parts.append(f"### {_('Labels')}: {group_title}")
                parts.append("")
                heading_level = "####"
            for view in group_views:
                req = view.requirement
                parts.append(f"{heading_level} {_requirement_heading(req, selected_fields)}")
                parts.append("")
                meta_rows: list[tuple[str, str, bool]] = []
                meta_rows.append((_('Requirement RID'), req.rid, False))
                if _should_render_field(selected_fields, "title"):
                    meta_rows.append((_('Title'), req.title or _('(no title)'), False))
                for field, label, use_code in _EXPORT_META_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    value = _meta_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    meta_rows.append((_(label), content, use_code))

                section_rows: list[tuple[str, str]] = []
                for field, label in _EXPORT_SECTION_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    value = _section_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    section_rows.append((_(label), content))

                for label, content, use_code in meta_rows:
                    value = _format_markdown_table_cell(content)
                    if use_code:
                        value = f"``{value}``"
                    parts.append(f"- **{label}:** {value}")
                if meta_rows:
                    parts.append("")

                for label, content in section_rows:
                    parts.append(f"**{label}**")
                    parts.append("")
                    parts.append(content)
                    parts.append("")

                if view.links and _should_render_field(selected_fields, "links"):
                    parts.append(f"**{_('Related requirements')}**")
                    for link in view.links:
                        label = link.rid
                        if link.exists:
                            label = f"[{link.rid}](#{link.rid})"
                        suffix: list[str] = []
                        if link.title:
                            suffix.append(link.title)
                        if not link.exists:
                            suffix.append(_('missing'))
                        if link.suspect:
                            suffix.append(_('suspect'))
                        if suffix:
                            parts.append(f"- {label} — {', '.join(suffix)}")
                        else:
                            parts.append(f"- {label}")
                    parts.append("")
    return "\n".join(parts).rstrip() + "\n"

def _escape_html(text: str) -> str:
    import html

    return html.escape(text)


_ATTACHMENT_LINK_RE = re.compile(r"!\[([^\]]*)\]\(attachment:([^)]+)\)")
_INLINE_FORMULA_RE = re.compile(r"\\\((.+?)\\\)|(?<!\\)\$(?!\$)(.+?)(?<!\\)\$")
_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def _split_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip() for cell in raw.split("|")]


def _latex_to_omml(latex: str) -> str | None:
    try:
        from latex2mathml.converter import convert as latex_to_mathml
        from mathml2omml import convert as mathml_to_omml
    except ImportError:
        return None
    try:
        mathml = latex_to_mathml(latex)
        return mathml_to_omml(mathml)
    except Exception:  # pragma: no cover - conversion failures
        return None


def _append_omml_run(paragraph: docx.text.paragraph.Paragraph, omml_xml: str) -> None:
    from docx.oxml import parse_xml

    if "xmlns:m=" not in omml_xml:
        if "<m:oMath" in omml_xml:
            omml_xml = omml_xml.replace(
                "<m:oMath",
                "<m:oMath xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\"",
                1,
            )
        else:
            omml_xml = (
                "<m:oMath xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\">"
                f"{omml_xml}</m:oMath>"
            )
    paragraph._p.append(parse_xml(omml_xml))


def _render_formula_run(
    paragraph: docx.text.paragraph.Paragraph,
    formula: str,
    *,
    formula_renderer: str,
) -> None:
    if formula_renderer in {"auto", "mathml"}:
        omml = _latex_to_omml(formula)
        if omml:
            _append_omml_run(paragraph, omml)
            return
    if formula_renderer in {"auto", "svg"}:
        image_bytes = _latex_to_svg_png(formula)
        if image_bytes:
            run = paragraph.add_run()
            run.add_picture(BytesIO(image_bytes))
            return
    if formula_renderer in {"auto", "png"}:
        image_bytes = _latex_to_png(formula)
        if image_bytes:
            run = paragraph.add_run()
            run.add_picture(BytesIO(image_bytes))
            return
    paragraph.add_run(formula)


def _looks_like_inline_formula(candidate: str) -> bool:
    stripped = candidate.strip()
    if not stripped:
        return False
    return any(ch.isalpha() for ch in stripped) or any(
        token in stripped for token in ("\\", "^", "_", "{", "}", "=", "+", "-", "*", "/")
    )


def _latex_to_png(latex: str) -> bytes | None:
    try:
        import matplotlib
        from matplotlib import pyplot as plt
    except ImportError:
        return None
    matplotlib.use("Agg", force=True)
    try:
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0.0, 0.0, f"${latex}$", fontsize=12)
        buffer = BytesIO()
        fig.savefig(
            buffer,
            format="png",
            bbox_inches="tight",
            pad_inches=0.1,
            transparent=True,
        )
        plt.close(fig)
        return buffer.getvalue()
    except Exception:  # pragma: no cover - rendering failures
        return None


def _latex_to_svg_png(latex: str) -> bytes | None:
    svg_bytes = _latex_to_svg(latex)
    if not svg_bytes:
        return None
    return _svg_to_png(svg_bytes)


def _latex_to_svg(latex: str) -> bytes | None:
    try:
        import matplotlib
        from matplotlib import pyplot as plt
    except ImportError:
        return None
    matplotlib.use("Agg", force=True)
    try:
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0.0, 0.0, f"${latex}$", fontsize=12)
        buffer = BytesIO()
        fig.savefig(
            buffer,
            format="svg",
            bbox_inches="tight",
            pad_inches=0.1,
            transparent=True,
        )
        plt.close(fig)
        return buffer.getvalue()
    except Exception:  # pragma: no cover - rendering failures
        return None


def _svg_to_png(svg_bytes: bytes) -> bytes | None:
    try:
        import cairosvg
    except ImportError:
        return None
    try:
        return cairosvg.svg2png(bytestring=svg_bytes)
    except Exception:  # pragma: no cover - rendering failures
        return None

def _build_markdown_renderer() -> markdown.Markdown:
    renderer = markdown.Markdown(
        extensions=[
            "markdown.extensions.tables",
            "markdown.extensions.fenced_code",
            "markdown.extensions.sane_lists",
            "markdown.extensions.nl2br",
        ],
        output_format="html5",
    )
    renderer.reset()
    return renderer


_MARKDOWN_RENDERER = _build_markdown_renderer()


def _render_markdown(text: str) -> str:
    renderer = _MARKDOWN_RENDERER
    renderer.reset()
    prepared = convert_markdown_math(text or "")
    markup = renderer.convert(prepared)
    return sanitize_html(markup)


def _attachment_markdown(text: str, *, requirement: Requirement) -> str:
    if "attachment:" not in text:
        return text
    attachment_map = {att.id: att.path for att in requirement.attachments}
    if not attachment_map:
        return text
    for attachment_id, path in attachment_map.items():
        text = text.replace(f"attachment:{attachment_id}", path)
    return text


def _html_markdown(value: str, *, requirement: Requirement) -> str:
    content = _attachment_markdown(value, requirement=requirement)
    return _render_markdown(content)


def _statement_preview(req: Requirement, *, max_chars: int) -> str:
    prepared = convert_markdown_math(req.statement or "")
    value = strip_markdown(prepared).strip()
    if len(value) <= max_chars:
        return value
    return value[: max(1, max_chars - 1)].rstrip() + "…"


def _build_preview_lookup(
    export: RequirementExport,
    *,
    max_statement_chars: int,
) -> dict[str, RequirementLinkPreview]:
    lookup: dict[str, RequirementLinkPreview] = {}
    for doc in export.documents:
        for view in doc.requirements:
            req = view.requirement
            lookup[req.rid] = RequirementLinkPreview(
                rid=req.rid,
                title=req.title or _('(no title)'),
                status=_meta_field_value(req, "status") or "",
                req_type=_meta_field_value(req, "type") or "",
                statement_preview=_statement_preview(req, max_chars=max_statement_chars),
            )
    return lookup


def _build_incoming_links(export: RequirementExport) -> dict[str, list[tuple[str, str]]]:
    incoming: dict[str, list[tuple[str, str]]] = {}
    for doc in export.documents:
        for view in doc.requirements:
            source = view.requirement
            source_label = source.title or _('(no title)')
            for link in view.links:
                bucket = incoming.setdefault(link.rid, [])
                bucket.append((source.rid, source_label))
    for rid in incoming:
        incoming[rid].sort(key=lambda item: item[0])
    return incoming


def _hierarchical_document_order(export: RequirementExport) -> list[DocumentExport]:
    selected = {doc.document.prefix: doc for doc in export.documents}
    children: dict[str | None, list[str]] = {}
    for prefix, payload in selected.items():
        parent = payload.document.parent
        if parent not in selected:
            parent = None
        children.setdefault(parent, []).append(prefix)
    for values in children.values():
        values.sort()

    ordered: list[DocumentExport] = []
    visited: set[str] = set()

    def walk(prefix: str) -> None:
        if prefix in visited:
            return
        visited.add(prefix)
        ordered.append(selected[prefix])
        for child in children.get(prefix, []):
            walk(child)

    for root in children.get(None, []):
        walk(root)
    for prefix in sorted(selected):
        walk(prefix)
    return ordered


def _json_for_html(payload: object) -> str:
    return json.dumps(payload, ensure_ascii=False).replace("</", "<\\/")


def render_requirements_html(
    export: RequirementExport,
    *,
    title: str | None = None,
    empty_field_placeholder: str | None = None,
    fields: Iterable[str] | None = None,
    group_by_labels: bool = False,
    unlabeled_group_title: str | None = None,
    label_group_mode: str = "per_label",
    colorize_label_backgrounds: bool = False,
    trace_mode: str = "flat",
    link_preview: bool = True,
    include_incoming_links: bool = False,
    max_preview_statement_chars: int = 220,
) -> str:
    """Render export data as standalone HTML."""
    selected_fields = _normalize_export_fields(fields)
    heading = title or _('Requirements export')
    parts: list[str] = [
        "<!DOCTYPE html>",
        "<html><head><meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1'>",
        f"<title>{_escape_html(heading)}</title>",
        "<style>html{font-size:16px;-webkit-text-size-adjust:100%;text-size-adjust:100%;}",
        "body{font-family:Arial,Helvetica,sans-serif;margin:24px;font-size:0.875rem;line-height:1.5;}",
        "h1{margin-top:0;font-size:1.5rem;}h2{font-size:1.25rem;}h3{font-size:1rem;}h4{font-size:0.875rem;margin-bottom:4px;}",
        "section.document{margin-bottom:32px;}",
        "article.requirement{border:1px solid #ddd;padding:16px;margin-bottom:16px;border-radius:8px;}",
        "article.requirement h3{margin-top:0;}article.requirement p{margin:0 0 8px;}",
        "table.meta-table{width:100%;border-collapse:collapse;margin:0 0 8px;}",
        "table.meta-table td{border:1px solid #ddd;padding:6px 8px;vertical-align:top;}",
        "table.meta-table tr:nth-child(even){background:#f2f2f2;}",
        "table.meta-table tr:nth-child(odd){background:#fff;}",
        "table.meta-table p{margin:0 0 8px;}",
        "table.meta-table .row-label{font-weight:bold;margin-bottom:4px;}",
        "dl.meta-list{margin:0 0 8px;padding:0;}",
        "dl.meta-list dt{font-weight:bold;margin:0;}",
        "dl.meta-list dd{margin:0 0 8px 0;}",
        "ul.links{margin:8px 0 0 16px;}",
        "ul.links li{margin-bottom:4px;}span.missing{color:#b00020;}span.suspect{color:#a35a00;}span.outside-export{color:#334f7d;}",
        "a.trace-link{position:relative;}",
        "#trace-preview-popover{position:fixed;z-index:9999;max-width:340px;padding:10px 12px;border:1px solid #C8CDD3;border-radius:8px;background:#fff;box-shadow:0 6px 18px rgba(0,0,0,.18);display:none;}",
        "#trace-preview-popover .rid{font-weight:700;margin-bottom:4px;}#trace-preview-popover .meta{color:#525252;font-size:.8125rem;margin-bottom:6px;}#trace-preview-popover .statement{font-size:.8125rem;line-height:1.35;white-space:pre-wrap;}",
        ".label-chip{display:inline-block;padding:2px 8px;border:1px solid #C8CDD3;border-radius:999px;font-size:0.8125rem;font-weight:600;line-height:1.2;background:#f5f6f8;color:#1f2328;margin:0 6px 4px 0;}",
        "</style>",
        "</head><body>",
        f"<h1>{_escape_html(heading)}</h1>",
        f"<p><em>{_escape_html(_('Generated at'))} {_render_generated_at(export)} {_escape_html(_('for documents'))}: {', '.join(export.selected_prefixes)}.</em></p>",
        f"<p><em>{_escape_html(_('Document revisions'))}: {_escape_html(_export_revisions_summary(export))}.</em></p>",
    ]
    if _should_render_field(selected_fields, "labels"):
        label_rows = _collect_used_label_rows(export)
        if label_rows:
            parts.append(f"<h2>{_escape_html(_('Labels'))}</h2>")
            parts.append("<table class='meta-table'><tbody>")
            for label, description, color in label_rows:
                if colorize_label_backgrounds:
                    label_html = _render_html_label_chip(label, color)
                else:
                    label_html = f"<strong>{_escape_html(label)}</strong>"
                parts.append(
                    "<tr>"
                    f"<td>{label_html}</td>"
                    f"<td>{_escape_html(description)}</td>"
                    "</tr>"
                )
            parts.append("</tbody></table>")

    palette = _label_palette(export) if colorize_label_backgrounds else {}
    preview_lookup = _build_preview_lookup(export, max_statement_chars=max_preview_statement_chars)
    rendered_rids = {view.requirement.rid for doc in export.documents for view in doc.requirements}
    incoming_links = _build_incoming_links(export) if include_incoming_links else {}
    document_iter = export.documents if trace_mode == "flat" else _hierarchical_document_order(export)

    for doc in document_iter:
        doc_prefix = _escape_html(doc.document.prefix)
        parts.append(f"<section class='document' id='doc-{doc_prefix}'>")
        parts.append(
            f"<h2>{_escape_html(doc.document.title)} (<code>{doc_prefix}</code>, {_escape_html(_document_revision_label(doc.document))})</h2>"
        )
        if group_by_labels:
            group_iter = _group_requirement_views_by_labels(
                doc.requirements,
                unlabeled_title=unlabeled_group_title or _('Without labels'),
                label_group_mode=label_group_mode,
            )
        else:
            group_iter = [("", list(doc.requirements))]

        for group_title, group_views in group_iter:
            if group_by_labels:
                parts.append(
                    f"<h3>{_escape_html(_('Labels'))}: {_escape_html(group_title)}</h3>"
                )
            for view in group_views:
                req = view.requirement
                parts.append(f"<article class='requirement' id='{_escape_html(req.rid)}'>")
                parts.append(
                    f"<h3>{_escape_html(_requirement_heading(req, selected_fields))}</h3>"
                )
                field_rows: list[tuple[str, str, bool]] = []
                field_rows.append((_('Requirement RID'), _escape_html(req.rid), True))
                if _should_render_field(selected_fields, "title"):
                    field_rows.append(
                        (_('Title'), _escape_html(req.title or _('(no title)')), True)
                    )
                for field, label, _use_code in _EXPORT_META_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    if field == "labels" and colorize_label_backgrounds:
                        labels = _normalized_labels(req)
                        if labels:
                            chips = "".join(
                                _render_html_label_chip(name, palette.get(name.casefold()))
                                for name in labels
                            )
                            field_rows.append((_(label), chips, False))
                            continue
                        content = _resolve_field_content(
                            None,
                            empty_field_placeholder=empty_field_placeholder,
                        )
                        if content is not None:
                            field_rows.append((_(label), _escape_html(content), True))
                        continue
                    value = _meta_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    field_rows.append((_(label), _escape_html(content), True))

                for field, label in _EXPORT_SECTION_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    value = _section_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    html_value = _html_markdown(content, requirement=req) or "<p></p>"
                    field_rows.append((_(label), html_value, False))

                if field_rows:
                    parts.append("<dl class='meta-list'>")
                    for label, value, _is_inline in field_rows:
                        parts.append(f"<dt>{_escape_html(label)}</dt>")
                        parts.append(f"<dd>{value}</dd>")
                    parts.append("</dl>")

                if view.links and _should_render_field(selected_fields, "links"):
                    parts.append(f"<h4>{_escape_html(_('Related requirements'))}</h4><ul class='links'>")
                    for link in view.links:
                        rid = _escape_html(link.rid)
                        title_value = link.title or preview_lookup.get(link.rid, RequirementLinkPreview(rid=link.rid, title="", status="", req_type="", statement_preview="")).title
                        title = _escape_html(title_value) if title_value else ""
                        classes: list[str] = []
                        if not link.exists:
                            classes.append("missing")
                        if link.suspect:
                            classes.append("suspect")
                        if link.exists and link.rid not in rendered_rids:
                            classes.append("outside-export")
                        cls_attr = f" class='{' '.join(classes)}'" if classes else ""
                        preview_attr = ""
                        if link_preview and link.rid in preview_lookup:
                            preview_attr = f" data-preview-id='{rid}'"
                        text = rid if not title else f"{rid} — {title}"
                        if link.exists and link.rid in rendered_rids:
                            parts.append(f"<li><a href='#{rid}' class='trace-link'{preview_attr}{cls_attr}>{text}</a></li>")
                        elif link.exists:
                            parts.append(f"<li><span{cls_attr}>{text} ({_escape_html(_('outside exported scope'))})</span></li>")
                        else:
                            text_parts = [rid]
                            if title:
                                text_parts.append(f"— {title}")
                            text_parts.append(f"({_('missing')})")
                            if link.suspect:
                                text_parts.append(f"({_('suspect')})")
                            parts.append(f"<li><span{cls_attr}>{' '.join(text_parts)}</span></li>")
                    parts.append("</ul>")

                if include_incoming_links and _should_render_field(selected_fields, "links"):
                    incoming = incoming_links.get(req.rid, [])
                    if incoming:
                        parts.append(f"<h4>{_escape_html(_('Linked from'))}</h4><ul class='links'>")
                        for source_rid, source_title in incoming:
                            source_rid_html = _escape_html(source_rid)
                            source_title_html = _escape_html(source_title)
                            source_text = source_rid_html if not source_title else f"{source_rid_html} — {source_title_html}"
                            preview_attr = ""
                            if link_preview and source_rid in preview_lookup:
                                preview_attr = f" data-preview-id='{source_rid_html}'"
                            if source_rid in rendered_rids:
                                parts.append(f"<li><a href='#{source_rid_html}' class='trace-link'{preview_attr}>{source_text}</a></li>")
                            else:
                                parts.append(f"<li><span class='outside-export'>{source_text} ({_escape_html(_('outside exported scope'))})</span></li>")
                        parts.append("</ul>")
                parts.append("</article>")
        parts.append("</section>")

    if link_preview and preview_lookup:
        preview_payload = {
            rid: {
                "rid": item.rid,
                "title": item.title,
                "status": item.status,
                "type": item.req_type,
                "statement": item.statement_preview,
                "exists": item.exists,
                "suspect": item.suspect,
            }
            for rid, item in preview_lookup.items()
        }
        parts.append("<div id='trace-preview-popover' role='tooltip' aria-hidden='true'></div>")
        parts.append(
            "<script type='application/json' id='trace-preview-data'>"
            f"{_json_for_html(preview_payload)}"
            "</script>"
        )
        script = """<script>(function(){
const payloadNode=document.getElementById('trace-preview-data');
if(!payloadNode){return;}
const payload=JSON.parse(payloadNode.textContent||'{}');
const popover=document.getElementById('trace-preview-popover');
if(!popover){return;}
const esc=(value)=>String(value??'').replace(/[&<>"']/g,(ch)=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[ch]));
const render=(item)=>{
  if(!item){return '';}
  const meta=[item.type,item.status].filter(Boolean).map(esc).join(' • ');
  return '<div class="rid">'+esc(item.rid)+'</div>'
    +'<div class="title">'+esc(item.title)+'</div>'
    +'<div class="meta">'+meta+'</div>'
    +'<div class="statement">'+esc(item.statement||'')+'</div>';
};
const place=(event)=>{
  const margin=12;
  let left=(event.clientX||0)+margin;
  let top=(event.clientY||0)+margin;
  const rect=popover.getBoundingClientRect();
  const maxLeft=window.innerWidth-rect.width-margin;
  const maxTop=window.innerHeight-rect.height-margin;
  left=Math.min(Math.max(margin,left),Math.max(margin,maxLeft));
  top=Math.min(Math.max(margin,top),Math.max(margin,maxTop));
  popover.style.left=left+'px';
  popover.style.top=top+'px';
};
const show=(event)=>{
  const target=event.currentTarget;
  const key=target.dataset.previewId;
  const item=payload[key];
  if(!item){return;}
  popover.innerHTML=render(item);
  popover.style.display='block';
  popover.setAttribute('aria-hidden','false');
  const rect=target.getBoundingClientRect();
  const synthetic={clientX:event.clientX||rect.left,clientY:event.clientY||rect.bottom};
  place(synthetic);
};
const hide=()=>{
  popover.style.display='none';
  popover.setAttribute('aria-hidden','true');
};
document.querySelectorAll('.trace-link[data-preview-id]').forEach((node)=>{
  node.addEventListener('mouseover',show);
  node.addEventListener('mousemove',place);
  node.addEventListener('mouseout',hide);
  node.addEventListener('focus',show);
  node.addEventListener('blur',hide);
});
})();</script>"""
        parts.append(script)
    parts.append("</body></html>")
    return "".join(parts)

def _iter_markdown_segments(
    text: str,
    *,
    attachment_map: dict[str, str],
) -> list[tuple[str, str]]:
    if "attachment:" not in text:
        return [("text", text)]
    segments: list[tuple[str, str]] = []
    start = 0
    for match in _ATTACHMENT_LINK_RE.finditer(text):
        if match.start() > start:
            segments.append(("text", text[start:match.start()]))
        attachment_id = match.group(2).strip()
        path = attachment_map.get(attachment_id)
        if path:
            segments.append(("image", path))
        else:
            segments.append(("text", match.group(1)))
        start = match.end()
    if start < len(text):
        segments.append(("text", text[start:]))
    return segments


def _docx_add_markdown(
    doc: docx.Document | docx.table._Cell,
    text: str,
    *,
    attachment_map: dict[str, str],
    base_path: Path,
    doc_prefix: str,
    image_width: float,
    formula_renderer: str,
    start_paragraph: docx.text.paragraph.Paragraph | None = None,
) -> None:
    normalized_text = normalize_escaped_newlines(text)
    segments = _iter_markdown_segments(normalized_text, attachment_map=attachment_map)
    first_paragraph = start_paragraph

    def _next_paragraph() -> docx.text.paragraph.Paragraph:
        nonlocal first_paragraph
        if first_paragraph is not None:
            paragraph = first_paragraph
            first_paragraph = None
            return paragraph
        return doc.add_paragraph()

    for kind, payload in segments:
        if kind == "image":
            image_path = base_path / doc_prefix / payload
            if image_path.exists():
                paragraph = _next_paragraph()
                run = paragraph.add_run()
                try:
                    run.add_picture(str(image_path), width=Inches(image_width))
                except (OSError, ValueError):  # pragma: no cover - invalid assets
                    _next_paragraph().add_run(strip_markdown(payload))
            else:
                _next_paragraph().add_run(strip_markdown(payload))
            continue
        lines = payload.splitlines()
        idx = 0
        while idx < len(lines):
            line = lines[idx]
            stripped = line.strip()
            if stripped.startswith("$$"):
                if stripped.endswith("$$") and len(stripped) > 4:
                    formula = stripped[2:-2].strip()
                    paragraph = _next_paragraph()
                    _render_formula_run(
                        paragraph,
                        formula,
                        formula_renderer=formula_renderer,
                    )
                    idx += 1
                    continue
                if stripped == "$$":
                    idx += 1
                    block_lines: list[str] = []
                    while idx < len(lines):
                        if lines[idx].strip() == "$$":
                            idx += 1
                            break
                        block_lines.append(lines[idx])
                        idx += 1
                    formula = "\n".join(block_lines).strip()
                    if formula:
                        paragraph = _next_paragraph()
                        _render_formula_run(
                            paragraph,
                            formula,
                            formula_renderer=formula_renderer,
                        )
                    continue
            if "|" in line and idx + 1 < len(lines) and _TABLE_SEPARATOR_RE.match(lines[idx + 1]):
                header_cells = _split_table_row(line)
                idx += 2
                table_rows: list[list[str]] = []
                while idx < len(lines):
                    row_line = lines[idx]
                    if "|" not in row_line:
                        break
                    row_cells = _split_table_row(row_line)
                    if row_cells:
                        table_rows.append(row_cells)
                    idx += 1
                if header_cells:
                    col_count = len(header_cells)
                    table = doc.add_table(rows=0, cols=col_count)
                    table.style = "Table Grid"
                    header_row = table.add_row().cells
                    for col_idx, cell in enumerate(header_cells):
                        header_row[col_idx].text = strip_markdown(cell)
                    for row in table_rows:
                        row_cells = table.add_row().cells
                        for col_idx, cell in enumerate(row[:col_count]):
                            row_cells[col_idx].text = strip_markdown(cell)
                continue
            if line.strip():
                paragraph = _next_paragraph()
                last_idx = 0
                for match in _INLINE_FORMULA_RE.finditer(line):
                    text_segment = line[last_idx:match.start()]
                    if text_segment:
                        paragraph.add_run(strip_markdown(text_segment))
                    formula = (match.group(1) or match.group(2) or "").strip()
                    if formula and _looks_like_inline_formula(formula):
                        _render_formula_run(
                            paragraph,
                            formula,
                            formula_renderer=formula_renderer,
                        )
                    elif match.group(0):
                        paragraph.add_run(strip_markdown(match.group(0)))
                    last_idx = match.end()
                tail = line[last_idx:]
                if tail:
                    paragraph.add_run(strip_markdown(tail))
            else:
                _next_paragraph()
            idx += 1


def _docx_needs_separate_label(content: str) -> bool:
    lines = content.splitlines()
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("$$"):
            return True
        if "|" in line and idx + 1 < len(lines) and _TABLE_SEPARATOR_RE.match(lines[idx + 1]):
            return True
    return False


def _docx_add_labeled_content(
    container: docx.Document | docx.table._Cell,
    label: str,
    content: str,
    *,
    attachment_map: dict[str, str],
    base_path: Path,
    doc_prefix: str,
    image_width: float,
    formula_renderer: str,
) -> None:
    label_text = f"{_(label)}:"
    normalized_content = content.strip("\n")
    start_paragraph = None
    if isinstance(container, docx.table._Cell):
        container.text = ""
        start_paragraph = container.paragraphs[0]
    if _docx_needs_separate_label(normalized_content):
        paragraph = start_paragraph or container.add_paragraph()
        label_run = paragraph.add_run(label_text)
        label_run.bold = True
        _docx_add_markdown(
            container,
            normalized_content,
            attachment_map=attachment_map,
            base_path=base_path,
            doc_prefix=doc_prefix,
            image_width=image_width,
            formula_renderer=formula_renderer,
        )
        return

    paragraph = start_paragraph or container.add_paragraph()
    label_run = paragraph.add_run(label_text)
    label_run.bold = True
    paragraph.add_run(" ")
    _docx_add_markdown(
        container,
        normalized_content,
        attachment_map=attachment_map,
        base_path=base_path,
        doc_prefix=doc_prefix,
        image_width=image_width,
        formula_renderer=formula_renderer,
        start_paragraph=paragraph,
    )


def _docx_style_run_as_label_chip(run: docx.text.run.Run, color: str | None) -> None:
    hex_color = _normalize_hex_color(color)
    if not hex_color:
        return
    run.font.color.rgb = RGBColor.from_string(_text_color_for_background(hex_color))
    run_properties = run._r.get_or_add_rPr()
    shading = run_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        run_properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)


def _docx_add_label_chips_line(
    cell: docx.table._Cell,
    *,
    label_text: str,
    labels: Sequence[str],
    palette: Mapping[str, str],
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if label_text:
        label_run = paragraph.add_run(f"{_(label_text)}: ")
        label_run.bold = True
    for index, label in enumerate(labels):
        run = paragraph.add_run(f"\u00A0{label}\u00A0")
        _docx_style_run_as_label_chip(run, palette.get(label.casefold()))
        if index + 1 < len(labels):
            paragraph.add_run(" ")


def _docx_apply_row_shading(row: docx.table._Row, *, fill: str) -> None:
    for cell in row.cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = cell_properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            cell_properties.append(shading)
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)


def render_requirements_docx(
    export: RequirementExport,
    *,
    title: str | None = None,
    formula_renderer: str = "auto",
    empty_field_placeholder: str | None = None,
    fields: Iterable[str] | None = None,
    group_by_labels: bool = False,
    unlabeled_group_title: str | None = None,
    label_group_mode: str = "per_label",
    colorize_label_backgrounds: bool = False,
) -> bytes:
    """Render export data as a DOCX document."""
    selected_fields = _normalize_export_fields(fields)
    heading = title or _('Requirements export')
    document = docx.Document()
    document.add_heading(heading, level=0)
    document.add_paragraph(
        f"{_('Generated at')} {_render_generated_at(export)} {_('for documents')}: {', '.join(export.selected_prefixes)}."
    )
    document.add_paragraph(
        f"{_('Document revisions')}: {_export_revisions_summary(export)}."
    )
    image_width = 5.5
    palette = _label_palette(export) if colorize_label_backgrounds else {}
    if _should_render_field(selected_fields, "labels"):
        label_rows = _collect_used_label_rows(export)
        if label_rows:
            document.add_heading(_('Labels'), level=1)
            label_table = document.add_table(rows=0, cols=2)
            label_table.style = "Table Grid"
            for row_index, (label, description, color) in enumerate(label_rows):
                row = label_table.add_row()
                if colorize_label_backgrounds:
                    _docx_add_label_chips_line(
                        row.cells[0],
                        label_text="",
                        labels=[label],
                        palette={label.casefold(): color or ""},
                    )
                else:
                    row.cells[0].text = label
                row.cells[1].text = description
                if row_index % 2 == 1:
                    _docx_apply_row_shading(row, fill="F2F2F2")
                else:
                    _docx_apply_row_shading(row, fill="FFFFFF")
            document.add_paragraph("")

    for doc_export in export.documents:
        document.add_heading(
            f"{doc_export.document.title} ({doc_export.document.prefix}, {_document_revision_label(doc_export.document)})",
            level=1,
        )
        if group_by_labels:
            group_iter = _group_requirement_views_by_labels(
                doc_export.requirements,
                unlabeled_title=unlabeled_group_title or _('Without labels'),
                label_group_mode=label_group_mode,
            )
        else:
            group_iter = [("", list(doc_export.requirements))]

        for group_title, group_views in group_iter:
            heading_level = 2
            if group_by_labels:
                document.add_heading(f"{_('Labels')}: {group_title}", level=2)
                heading_level = 3
            for view in group_views:
                req = view.requirement
                document.add_heading(_requirement_heading(req, selected_fields), level=heading_level)
                field_rows: list[tuple[str, str | None, tuple[str, ...] | None]] = []
                field_rows.append(("Requirement RID", req.rid, None))
                if _should_render_field(selected_fields, "title"):
                    field_rows.append(("Title", req.title or _('(no title)'), None))
                for field, label, _use_code in _EXPORT_META_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    if field == "labels" and colorize_label_backgrounds:
                        labels = _normalized_labels(req)
                        if labels:
                            field_rows.append((label, None, labels))
                            continue
                        content = _resolve_field_content(
                            None,
                            empty_field_placeholder=empty_field_placeholder,
                        )
                        if content is not None:
                            field_rows.append((label, content, None))
                        continue
                    value = _meta_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    field_rows.append((label, content, None))

                attachment_map = {att.id: att.path for att in req.attachments}
                for field, label in _EXPORT_SECTION_FIELDS:
                    if not _should_render_field(selected_fields, field):
                        continue
                    value = _section_field_value(req, field)
                    content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                    if content is None:
                        continue
                    field_rows.append((label, content, None))
                if field_rows:
                    table = document.add_table(rows=0, cols=1)
                    table.style = "Table Grid"
                    for row_index, (label, content, label_chips) in enumerate(field_rows):
                        row = table.add_row()
                        if label_chips is not None:
                            _docx_add_label_chips_line(
                                row.cells[0],
                                label_text=label,
                                labels=label_chips,
                                palette=palette,
                            )
                        else:
                            _docx_add_labeled_content(
                                row.cells[0],
                                label,
                                content,
                                attachment_map=attachment_map,
                                base_path=export.base_path,
                                doc_prefix=req.doc_prefix,
                                image_width=image_width,
                                formula_renderer=formula_renderer,
                            )
                        if row_index % 2 == 1:
                            _docx_apply_row_shading(row, fill="F2F2F2")
                        else:
                            _docx_apply_row_shading(row, fill="FFFFFF")
                document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _ensure_stylesheet() -> StyleSheet1:
    styles = getSampleStyleSheet()
    if "RequirementHeading" not in styles:
        styles.add(
            ParagraphStyle(
                "RequirementHeading",
                parent=styles["Heading3"],
                spaceBefore=12,
                spaceAfter=6,
            )
        )
    if "SectionHeading" not in styles:
        styles.add(
            ParagraphStyle(
                "SectionHeading",
                parent=styles["Heading4"],
                fontSize=11,
                leading=14,
                spaceBefore=6,
                spaceAfter=4,
            )
        )
    if "MetaValue" not in styles:
        styles.add(
            ParagraphStyle(
                "MetaValue",
                parent=styles["BodyText"],
                spaceBefore=0,
                spaceAfter=0,
            )
        )
    return styles


def _pdf_text(value: str) -> str:
    return xml_escape(value).replace("\n", "<br/>")


def render_requirements_pdf(
    export: RequirementExport,
    *,
    title: str | None = None,
    empty_field_placeholder: str | None = None,
    fields: Iterable[str] | None = None,
) -> bytes:
    """Render export data as a PDF document."""
    selected_fields = _normalize_export_fields(fields)
    buffer = BytesIO()
    heading = title or _('Requirements export')
    styles = _ensure_stylesheet()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=heading,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    story: list = []
    story.append(Paragraph(xml_escape(heading), styles["Title"]))
    story.append(
        Paragraph(
            xml_escape(
                f"{_('Generated at')} {_render_generated_at(export)} {_('for documents')}: {', '.join(export.selected_prefixes)}."
            ),
            styles["BodyText"],
        )
    )
    story.append(
        Paragraph(
            xml_escape(f"{_('Document revisions')}: {_export_revisions_summary(export)}."),
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))

    for doc_export in export.documents:
        story.append(
            Paragraph(
                xml_escape(
                    f"{doc_export.document.title} ({doc_export.document.prefix}, {_document_revision_label(doc_export.document)})"
                ),
                styles["Heading2"],
            )
        )
        story.append(Spacer(1, 6))
        for view in doc_export.requirements:
            req = view.requirement
            story.append(
                Paragraph(
                    f"<a name='{xml_escape(req.rid)}'/><b>{xml_escape(_requirement_heading(req, selected_fields))}</b>",
                    styles["RequirementHeading"],
                )
            )
            data: list[list[str]] = []
            for field, label, _use_code in _EXPORT_META_FIELDS:
                if not _should_render_field(selected_fields, field):
                    continue
                value = _meta_field_value(req, field)
                content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                if content is None:
                    continue
                data.append([xml_escape(_(label)), _pdf_text(content)])
            if data:
                table = Table(data, colWidths=[40 * mm, 120 * mm])
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 6))

            for field, label in _EXPORT_SECTION_FIELDS:
                if not _should_render_field(selected_fields, field):
                    continue
                value = _section_field_value(req, field)
                content = _resolve_field_content(value, empty_field_placeholder=empty_field_placeholder)
                if content is None:
                    continue
                story.append(Paragraph(xml_escape(_(label)), styles["SectionHeading"]))
                story.append(Paragraph(_pdf_text(content), styles["BodyText"]))

            if view.links and _should_render_field(selected_fields, "links"):
                items = []
                for link in view.links:
                    label = xml_escape(link.rid)
                    if link.exists:
                        text = label
                        if link.title:
                            text += f" — {xml_escape(link.title)}"
                        if link.suspect:
                            text += f" ({_('suspect')})"
                        items.append(
                            ListItem(
                                Paragraph(
                                    f"<link href='#{label}' color='blue'>{text}</link>",
                                    styles["BodyText"],
                                )
                            )
                        )
                    else:
                        text = label
                        if link.title:
                            text += f" — {xml_escape(link.title)}"
                        text += f" ({_('missing')})"
                        if link.suspect:
                            text += f" ({_('suspect')})"
                        items.append(ListItem(Paragraph(text, styles["BodyText"])))
                story.append(Paragraph(xml_escape(_('Related requirements')), styles["SectionHeading"]))
                story.append(ListFlowable(items, bulletType="bullet"))
            story.append(Spacer(1, 12))
        story.append(Spacer(1, 12))

    doc.build(story)
    return buffer.getvalue()
